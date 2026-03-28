#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将Word文档转换为Markdown格式"""

from docx import Document
import os

def convert_docx_to_md(docx_path, md_path):
    """将docx文件转换为markdown文件"""
    
    # 读取Word文档
    doc = Document(docx_path)
    
    md_lines = []
    
    # 处理所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            md_lines.append("")
            continue
        
        # 判断段落样式
        style_name = para.style.name if para.style else ""
        
        # 处理标题
        if style_name.startswith("Heading 1") or "标题 1" in style_name:
            md_lines.append(f"# {text}")
        elif style_name.startswith("Heading 2") or "标题 2" in style_name:
            md_lines.append(f"## {text}")
        elif style_name.startswith("Heading 3") or "标题 3" in style_name:
            md_lines.append(f"### {text}")
        elif style_name.startswith("Heading") or "标题" in style_name:
            md_lines.append(f"#### {text}")
        # 处理列表
        elif para.paragraph_format.left_indent is not None and para.paragraph_format.left_indent > 0:
            indent = int(para.paragraph_format.left_indent / 914400)  # 转换为缩进级别
            bullet = "  " * indent + "- " if indent > 0 else "- "
            md_lines.append(f"{bullet}{text}")
        else:
            md_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        md_lines.append("")
        rows = table.rows
        if rows:
            # 表头
            header_cells = [cell.text.strip() for cell in rows[0].cells]
            md_lines.append("| " + " | ".join(header_cells) + " |")
            md_lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
            
            # 表体
            for row in rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")
    
    # 写入Markdown文件
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))
    
    print(f"✓ 转换完成！")
    print(f"  源文件: {docx_path}")
    print(f"  目标文件: {md_path}")

if __name__ == "__main__":
    # 文件路径
    docx_file = "AI网文写作系统白皮书v1.0.docx"
    md_file = "AI网文写作系统白皮书.md"
    
    if os.path.exists(docx_file):
        convert_docx_to_md(docx_file, md_file)
    else:
        print(f"错误: 找不到文件 {docx_file}")
